from flask import send_file
from sqlalchemy.orm.exc import NoResultFound
from sqlalchemy.orm import joinedload
from sqlalchemy import or_
from .. import db
from app.models import Content, Job
from bson import ObjectId
import json, requests, os, random, time
from openai import OpenAI
from bs4 import BeautifulSoup
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from docx import Document
from html2docx import html2docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from logging import getLogger
# from .chat import do_chat

# Initialize Logger
logger = getLogger(__name__)

API_KEY = os.getenv("OPENROUTER_API_KEY")
OPENROUTER_API = os.getenv("OPENROUTER_API")

# Initialize OpenAI client
openai_client = OpenAI(base_url=OPENROUTER_API, api_key=API_KEY)

def set_paragraph_rtl(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.rtl = True

def set_paragraph_rtl_and_center(paragraph):
    paragraph.paragraph_format.rtl = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def html_to_docx(html_string):
    try:
        doc = Document()
        section = doc.sections[0]
        section.right_to_left = True

        soup = BeautifulSoup(html_string, "html.parser")
        
        for elem in soup.descendants:
            paragraph = None
            if elem.name and elem.name.startswith('h'):
                level = int(elem.name[1])
                paragraph = doc.add_heading(elem.get_text(), level=level)
            elif elem.name == "p":
                paragraph = doc.add_paragraph(elem.get_text())
            elif elem.name == "ul":
                for li in elem.find_all('li'):
                    paragraph = doc.add_paragraph(f'- {li.get_text()}', style='ListBullet')
            elif elem.name == "ol":
                for li in elem.find_all('li'):
                    paragraph = doc.add_paragraph(li.get_text(), style='ListNumber')
            elif elem.name == "a":
                tag_text = elem.get_text()
                href = elem.get('href', '#')
                paragraph = doc.add_paragraph(f'{tag_text} ({href})')
            elif elem.name == "b":
                run = doc.add_paragraph().add_run(elem.get_text())
                run.bold = True
                paragraph = run.paragraph
            elif elem.name == "i":
                run = doc.add_paragraph().add_run(elem.get_text())
                run.italic = True
                paragraph = run.paragraph
            elif elem.name == "img":
                img_url = elem.get('src')
                try:
                    response = requests.get(img_url)
                    img_bytes = response.content
                    paragraph = doc.add_paragraph()
                    paragraph.add_run().add_picture(BytesIO(img_bytes), width=Inches(6))
                    set_paragraph_rtl_and_center(paragraph)
                except Exception as e:
                    logger.error(f"Error adding image: {e}")
            if paragraph:
                set_paragraph_rtl(paragraph)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        logger.error(f"Error converting HTML to DOCX: {e}")
        raise e
    
def suggest_search_query(input):
    try:
        llm_type="gpt-4o-mini"
        user_prompt = (
            f"Write a SEO-optimized query to image search for this title {input}."
            f"This image is going to use as part of an article"
        )
        assistant_prompt = (
            "You write SEO-optimized query for earch"

        )

        messages = [
            {"role": "assistant", "content": assistant_prompt},
            {"role": "user", "content": user_prompt},
        ]

        response = openai_client.chat.completions.create(
            model=llm_type, messages=messages, temperature=0.8
        )
        result = response.choices[0].message.content


        return result
    except Exception as e:
        logger.error(f"Error suggesting query for search of '{input}': {e}")
        return []


def suggest_one_image(selected_outline):
    return get_search_images(selected_outline, 1)
        

def save_html_to_docx(html_string, file_path):
    buffer = html_to_docx(html_string)
    with open(file_path, 'wb') as f:
        f.write(buffer.read())

def suggest_outlines(user_title, lang, llm_type="gpt-4o"):
    try:
        llm_type="gpt-4o"
        user_prompt = f"The content language is {lang}\n"
        user_prompt += f"Use numbering for list. \n"
        random.seed(int(time.time()))
        rand = random.randint(4, 6)
        user_prompt += (
            f"Generate two detailed outline for a blog post on the following topic: {user_title}.\n"
            f"The outline should include an introduction, {rand} main headlines (h1) and 2 to 3 sub-headlines (h2), and a conclusion.\n"
            f"Ensure each section flows logically and covers the topic comprehensively.\n"
            f"Template:"
            f"<list>"
            f"<block>"
            f"<h1> .. </h1>"
            f"<h2> .. </h2>"
            f"<h2> .. </h2>"
            f"<h2> .. </h2>"
            f"</block>"
            f"<block>"
            f"<h1> .. </h1>"
            f"<h2> .. </h2>"
            f"<h2> .. </h2>"
            f"<h2> .. </h2>"
            f"</block>"
            f"...\n"
            f"</list>"
        )

        assistant_prompt = (
            f"You write SEO-optimized blog posts' titles. Ensure correct grammar.\n"
            f"Only return the outlines and h1 and h2 HTML tags.\n"
        )

        messages = [
            {"role": "assistant", "content": assistant_prompt},
            {"role": "user", "content": user_prompt},
        ]

        response = openai_client.chat.completions.create(
            model=llm_type, messages=messages, temperature=0.8
        )
        result = response.choices[0].message.content

        soup = BeautifulSoup(result, "html.parser")
        lists = []

        for list_tag in soup.find_all("list"):
            list_dict = []
            for block in list_tag.find_all("block"):
                h1 = block.find("h1").get_text() if block.find("h1") else None
                h2_dict = [h2_tag.get_text() for h2_tag in block.find_all("h2")]
                list_dict.append({"head": h1, "subs": h2_dict})
            lists.append(list_dict)

        logger.info(f"Generated outlines for topic: {user_title}")
        if len(lists) < 2:
            logger.info(f"Unsuccessful generating outline for {user_title}, retrying")
            return suggest_outlines(user_title, lang, llm_type)
        return lists
    except Exception as e:
        logger.error(f"Error suggesting outlines for topic '{user_title}': {e}")
        raise e

def suggest_titles(topic, lang, llm_type="gpt-4o"):
    try:
        llm_type="gpt-4o"
        user_prompt = (
            f"Generate 6 catchy and informative title for a blog post about {topic}. "
            f"The title should be engaging and appealing to the general audience and "
            f"should reflect a professional tone. The content language is {lang}\n"
        )
        assistant_prompt = (
            "You write SEO-optimized titles for articles. Ensure correct grammar. "
            "Only return the in <h1> tags.\n"
            "<h1> .. </h1>\n"
            "<h1> .. </h1>\n"
            "<h1> .. </h1>\n"
        )

        messages = [
            {"role": "assistant", "content": assistant_prompt},
            {"role": "user", "content": user_prompt},
        ]

        response = openai_client.chat.completions.create(
            model=llm_type, messages=messages, temperature=0.8,
        )
        result = response.choices[0].message.content

        soup = BeautifulSoup(result, "html.parser")
        headlines_elements = soup.find_all("h1")
        titles = [headline_element.get_text() for headline_element in headlines_elements]

        logger.info(f"Generated titles for topic: {topic}")
        return titles
    except Exception as e:
        logger.error(f"Error suggesting titles for topic '{topic}': {e}")
        return []

def search_resources(topic):
    try:
        base_url = os.getenv("WEB_SEARCH_API")
        search = f"{base_url}?q={topic}&format=json"
        response = requests.get(search)
        result = response.json()

        logger.info(f"Searched resources for topic: {topic}")
        return result["results"][:10]
    except Exception as e:
        logger.error(f"Error searching resources for topic '{topic}': {e}")
        return []

def get_search_images(topic, num):
    try:
        query = suggest_search_query(topic)
        base_url = os.getenv("WEB_SEARCH_API")
        #search = f"{base_url}?q={topic}&format=json&categories=images"
        search = f"{base_url}?q={query}&format=json&categories=images&engine=!bi"        
        response = requests.get(search)
        result = response.json()

        logger.info(f"Searched images for topic: {topic}")
        return result["results"][:num]
    except Exception as e:
        logger.error(f"Error searching images for topic '{topic}': {e}")
        return []

def get_user_contents(user, search_query="", sort_order="desc", page=1, per_page=5):
    try:
        query = Content.query.filter_by(author_id=user.id)
        if search_query:
            query = query.filter(Content.system_title.like(f"%{search_query}%"))

        if sort_order == "asc":
            query = query.order_by(Content.timestamp.asc())
        else:
            query = query.order_by(Content.timestamp.desc())

        logger.info(f"Fetched contents for user {user.id} with search query '{search_query}' and sort order '{sort_order}'")
        return query.paginate(page, per_page, False)
    except Exception as e:
        logger.error(f"Error fetching user contents for user {user.id}: {e}")
        return []

def update_content(content_id, user_input):
    try:
        content = Content.query.get_or_404(content_id)

        body = user_input.get("body")
        content.body = body

        content.user_input = json.dumps(user_input)
        content.content_type = user_input.get("content_type")
        content.system_title = user_input.get("system_title")

        db.session.commit()
        
        logger.info(f"Updated content ID {content_id}")
        return content
    except NoResultFound:
        logger.warning(f"Content not found for ID {content_id}")
        raise ValueError("Content not found")
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error updating content ID {content_id}: {e}")
        raise e

def update_article_pro(content_id, body):
    try:
        content = Content.query.get_or_404(content_id)
        content.body = body
        db.session.commit()
        
        logger.info(f"Updated pro article content ID {content_id}")
        return content
    except NoResultFound:
        logger.warning(f"Pro article content not found for ID {content_id}")
        raise ValueError("Content not found")
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error updating pro article content ID {content_id}: {e}")
        raise e

def delete_article_pro(content_id):
    try:
        content = Content.query.get_or_404(content_id)
        db.session.delete(content)
        db.session.commit()

        logger.info(f"Deleted pro article content ID {content_id}")
        return content
    except NoResultFound:
        logger.warning(f"Pro article content not found for ID {content_id}")
        raise ValueError("Content not found")
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error deleting pro article content ID {content_id}: {e}")
        raise e


def create_content(user_input, author, llm=None):
    # Extract body from user input
    body = user_input.get("body")
    if user_input.get('llm') and not llm:
        llm = user_input.get('llm')
    if user_input.get('language_model') and not llm:
        llm = user_input.get('language_model')
    # Store body in MongoDB
    
    content_type = user_input.get("content_type")

    # Create SQL content instance
    content = Content(
        user_input=json.dumps(user_input),
        author=author,
        content_type=content_type,
        llm=llm,
        body=body
    )
    db.session.add(content)
    db.session.commit()
    return content


def create_job_record(job_id, content):
    try:
        job_record = Job(job_status="PENDING", job_id=job_id, content=content)
        db.session.add(job_record)
        db.session.commit()
        
        logger.info(f"Created job record with job ID {job_id}")
        return job_record
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error creating job record with job ID {job_id}: {e}")
        raise e

def get_job_by_id(id):
    try:
        if not id.isdigit():
            return None
        job = Job.query.filter_by(id=id).first()
        if job:
            logger.info(f"Fetched job ID: {id}")
        return job
    except Exception as e:
        logger.error(f"Error fetching job by ID {id}: {e}")
        return None

def get_job_by_cid(job_id):
    try:
        job = Job.query.filter_by(job_id=job_id).first()
        if job:
            logger.info(f"Fetched job by CID: {job_id}")
        return job
    except Exception as e:
        logger.error(f"Error fetching job by CID {job_id}: {e}")
        return None

def get_content_info(content_id):
    try:
        content = Content.query.get_or_404(content_id)
        logger.info(f"Fetched content info for ID: {content_id}")
        return content
    except Exception as e:
        logger.error(f"Error fetching content info by ID {content_id}: {e}")
        return None

def get_content_by_id(content_id):
    try:
        content = db.session.query(Content).options(joinedload(Content.job)).get(content_id)
        return content
    except Exception as e:
        logger.error(f"Error fetching content by ID {content_id}: {e}")
        return None
    
def is_content_has_feedback(content_id):
    content = Content.query.get_or_404(content_id)
    feedback = content.get_input('feedback')
    if feedback: 
        return feedback
    else: return ""
        
